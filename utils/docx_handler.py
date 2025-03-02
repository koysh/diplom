from docx import Document
import io
import logging

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Извлекает текст из DOCX-файла."""
    try:
        file_stream = io.BytesIO(file_bytes)  # Преобразуем в поток
        doc = Document(file_stream)  # Загружаем документ

        text_list = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_list.append(cell.text.strip())

        text = "\n".join(text_list)
        logger.info(f"Извлеченный текст из DOCX: {text}")  # Логируем извлеченный текст
        return text if text else "Файл пуст или не содержит текста"
    except Exception as e:
        logger.error(f"Ошибка при обработке DOCX: {e}")
        return f"Ошибка при обработке DOCX: {e}"